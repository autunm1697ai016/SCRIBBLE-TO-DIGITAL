import csv
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def build_txt_bytes(clean_notes: str, tasks: list) -> bytes:
    """Build TXT file bytes"""
    content = f"Clean Notes:\n{clean_notes}\n\nTasks:\n" + "\n".join(f"- {task}" for task in tasks)
    return content.encode('utf-8')

def build_csv_bytes(tasks: list) -> bytes:
    """Build CSV file bytes"""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Task"])
    for task in tasks:
        writer.writerow([task])
    return output.getvalue().encode('utf-8')

def build_pdf_bytes(clean_notes: str, tasks: list) -> bytes:
    """Build PDF file bytes"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    story = []
    story.append(Paragraph("Clean Notes", styles['Heading1']))
    story.append(Paragraph(clean_notes, styles['Normal']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Tasks", styles['Heading1']))
    for task in tasks:
        story.append(Paragraph(f"• {task}", styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def build_docx_bytes(clean_notes: str, tasks: list) -> bytes:
    """Build DOCX file bytes"""
    doc = Document()
    doc.add_heading('Clean Notes', 0)
    doc.add_paragraph(clean_notes)
    doc.add_heading('Tasks', 0)
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()